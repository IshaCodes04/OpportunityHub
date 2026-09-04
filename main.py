import logging
from selenium import webdriver  # type: ignore
from selenium.webdriver.common.by import By  # type: ignore
from selenium.webdriver.support.ui import WebDriverWait  # type: ignore
from selenium.webdriver.support import expected_conditions as EC  # type: ignore
from bs4 import BeautifulSoup  # type: ignore
import pandas as pd
import random
import time
import os
from datetime import datetime, timezone
import json
import re
import csv
import argparse
import html
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from urllib.parse import quote_plus
from urllib.request import Request, urlopen
from urllib.error import HTTPError, URLError

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger("opportunityhub")

# Target roles for scraping
TARGET_ROLES = [
    "full stack developer",
    "full stack engineer",
    "mern stack developer",
    "software engineer",
    "software developer",
    "junior software engineer",
    "sde",
    "sde-1",
    "software development engineer",
    "junior full stack developer",
    "associate software engineer",
    "graduate software engineer",
    "trainee software engineer",
    "generative ai engineer",
    "genai engineer",
    "genai developer",
    "ai engineer",
    "ai application developer",
    "ai software engineer",
    "ai developer",
    "full stack ai engineer",
    "ai application engineer",
    "llm application developer",
    "llm engineer",
    "rag engineer",
    "generative ai developer",
    "junior ai engineer",
    "junior genai engineer",
]

ALLOWED_ROLE_PATTERNS = [
    r"\bfull\s*stack\s+(?:developer|engineer)\b",
    r"\bmern\s+stack\s+developer\b",
    r"\bsoftware\s+(?:engineer|developer)\b",
    r"\bsde(?:\s*[- ]?1)?\b(?!\s*(?:ii|2|iii|iv|4|v|5)\b)",
    r"\bsoftware\s+development\s+engineer\b",
    r"\bgenerative\s+ai\s+engineer\b",
    r"\bgen\s*ai\s+engineer\b",
    r"\bgen\s*ai\s+developer\b",
    r"\bai\s+engineer\b",
    r"\bai\s+application\s+developer\b",
    r"\bai\s+software\s+engineer\b",
    r"\bai\s+developer\b",
    r"\bfull\s+stack\s+ai\s+engineer\b",
    r"\bai\s+application\s+engineer\b",
    r"\bllm\s+application\s+developer\b",
    r"\bllm\s+engineer\b",
    r"\brag\s+engineer\b",
    r"\bgenerative\s+ai\s+developer\b",
]

DISALLOWED_SENIORITY_PATTERNS = [
    r"\b(?:senior|sr|sr\.|lead|principal|staff|manager|director|head|architect)\b",
]

# Target locations — only jobs from these cities will be captured
TARGET_LOCATIONS = [
    # ── Delhi NCR ──
    "new delhi",
    "delhi",
    "gurgaon",
    "gurugram",
    "noida",
    "dwarka",
    "greater noida",
]

TARGET_COMPANIES = [
    {"ats": "lever", "slug": "zeta", "name": "Zeta"},
    {"ats": "greenhouse", "slug": "turing", "name": "Turing"},
    {"ats": "lever", "slug": "icertis", "name": "Icertis"},
    {"ats": "lever", "slug": "locus.co", "name": "Locus"},
    {"ats": "greenhouse", "slug": "netradyne", "name": "Netradyne"},
    {"ats": "lever", "slug": "curefit", "name": "Cure.fit"},
    {"ats": "greenhouse", "slug": "inmobi", "name": "InMobi"},
    {"ats": "lever", "slug": "hevodata", "name": "Hevo Data"},
    {"ats": "greenhouse", "slug": "groww", "name": "Groww"},
    {"ats": "lever", "slug": "paytm", "name": "Paytm"},
    {"ats": "lever", "slug": "meesho", "name": "Meesho"},
    {"ats": "greenhouse", "slug": "zenoti", "name": "Zenoti"},
    {"ats": "ashby", "slug": "plotlineso", "name": "Plotline"},
    {"ats": "ashby", "slug": "playpowerlabs", "name": "Playpower Labs"},
    {"ats": "lever", "slug": "upstox", "name": "Upstox"},
    {"ats": "greenhouse", "slug": "nanonets", "name": "Nanonets"},
    {"ats": "greenhouse", "slug": "aidashinc", "name": "AiDASH"},
    {"ats": "greenhouse", "slug": "ocrolusinc", "name": "Ocrolus"},
    {"ats": "greenhouse", "slug": "saucelabs", "name": "Sauce Labs"},
    {"ats": "greenhouse", "slug": "gravitonresearchcapital", "name": "Graviton Research Capital"},
    {"ats": "greenhouse", "slug": "nksecuritiesresearch", "name": "NK Securities Research"},
    {"ats": "greenhouse", "slug": "pay2dc", "name": "PayPay India"},
    {"ats": "ashby", "slug": "commure", "name": "Commure"},
    {"ats": "ashby", "slug": "sentilink", "name": "SentiLink"},
    {"ats": "ashby", "slug": "better-mortgage", "name": "Better Mortgage"},
    {"ats": "lever", "slug": "klearnow", "name": "KlearNow.ai"},
    {"ats": "lever", "slug": "thinkahead", "name": "AHEAD"},
    {"ats": "lever", "slug": "resilinc", "name": "Resilinc"},
    {"ats": "greenhouse", "slug": "swiggy", "name": "Swiggy"},
    {"ats": "greenhouse", "slug": "razorpay", "name": "Razorpay"},
    {"ats": "greenhouse", "slug": "freshworks", "name": "Freshworks"},
    {"ats": "lever", "slug": "vedantu", "name": "Vedantu"},
    {"ats": "greenhouse", "slug": "unacademy", "name": "Unacademy"},
    {"ats": "ashby", "slug": "cred", "name": "CRED"},
    {"ats": "lever", "slug": "dunzo", "name": "Dunzo"},
    {"ats": "greenhouse", "slug": "phonepe", "name": "PhonePe"},
    {"ats": "greenhouse", "slug": "flipkart", "name": "Flipkart"},
    {"ats": "greenhouse", "slug": "hiddenlayer", "name": "Hidden Layer"},
    {"ats": "lever", "slug": "akkio", "name": "Akkio"},
    {"ats": "ashby", "slug": "humartificial", "name": "Humartificial"},
    {"ats": "greenhouse", "slug": "gpt4all", "name": "GPT4All"},
    {"ats": "greenhouse", "slug": "thoughtworks", "name": "ThoughtWorks"},
    {"ats": "greenhouse", "slug": "harelabs", "name": "Hare Labs"},
    {"ats": "ashby", "slug": "appliedai", "name": "Applied AI"},
    {"ats": "lever", "slug": "almabetter", "name": "Alma Better"},
    {"ats": "lever", "slug": "unstop", "name": "Unstop"},
]

# ─────────────────────────── Time helpers ────────────────────────────

from datetime import timedelta

# Indian Standard Time = UTC + 5:30
IST = timezone(timedelta(hours=5, minutes=30))


def _ist_now() -> datetime:
    """Returns current time in IST."""
    return datetime.now(IST)


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def ist_timestamp_str(dt: datetime | None = None) -> str:
    """Returns IST timestamp string like: 2026-04-23 09:56 PM IST"""
    dt = dt or _ist_now()
    if dt.tzinfo != IST:
        dt = dt.astimezone(IST)
    return dt.strftime("%Y-%m-%d %I:%M %p IST")


def utc_timestamp_str(dt: datetime | None = None) -> str:
    """Returns UTC timestamp string like: 2026-04-23 10:36 UTC"""
    dt = dt or _utc_now()
    if dt.tzinfo != timezone.utc:
        dt = dt.astimezone(timezone.utc)
    return dt.strftime("%Y-%m-%d %I:%M %p UTC")


def get_target_date() -> str:
    """Returns today's date in IST (YYYY-MM-DD), overridable via TARGET_DATE env var."""
    env = os.getenv("TARGET_DATE")
    if env and len(env) == 10 and env[4] == "-" and env[7] == "-":
        return env
    return _ist_now().strftime("%Y-%m-%d")


def day_label_from_target_date(target_date: str) -> str:
    year, month, day = target_date.split("-")
    dt = datetime(int(year), int(month), int(day), tzinfo=IST)
    return dt.strftime("%B %#d, %Y") if os.name == "nt" else dt.strftime("%B %-d, %Y")


# ─────────────────────────── Filename helpers ────────────────────────

def output_hourly_csv_filename(target_date: str, now: datetime) -> str:
    """e.g. opportunities_2026_04_23_09_56_PM.csv (IST)"""
    ist = now.astimezone(IST)
    return f"opportunities_{target_date.replace('-', '_')}_{ist.strftime('%I_%M_%p')}.csv"


def output_hourly_txt_filename(target_date: str, now: datetime) -> str:
    """e.g. opportunities_2026_04_23_09_56_PM.txt (IST)"""
    ist = now.astimezone(IST)
    return f"opportunities_{target_date.replace('-', '_')}_{ist.strftime('%I_%M_%p')}.txt"


def role_to_slug(role: str) -> str:
    """Convert role name to folder/filename safe slug.
    e.g. 'Full Stack Developer' -> 'full_stack_developer'
    """
    return re.sub(r'[^a-z0-9]+', '_', role.lower()).strip('_')


def output_role_csv_path(role: str, target_date: str, now: datetime) -> str:
    """e.g. jobs/full_stack_developer/full_stack_developer_2026_04_23_09_56_PM.csv (IST)"""
    slug = role_to_slug(role)
    folder = os.path.join("jobs", slug)
    os.makedirs(folder, exist_ok=True)
    ist = now.astimezone(IST)
    filename = f"{slug}_{target_date.replace('-', '_')}_{ist.strftime('%I_%M_%p')}.csv"
    return os.path.join(folder, filename)


def output_role_txt_path(role: str, target_date: str, now: datetime) -> str:
    """e.g. jobs/full_stack_developer/full_stack_developer_2026_04_23_09_56_PM.txt (IST)"""
    slug = role_to_slug(role)
    folder = os.path.join("jobs", slug)
    os.makedirs(folder, exist_ok=True)
    ist = now.astimezone(IST)
    filename = f"{slug}_{target_date.replace('-', '_')}_{ist.strftime('%I_%M_%p')}.txt"
    return os.path.join(folder, filename)


# ─────────────────────── seen_jobs.json helpers ───────────────────────
# Format: { url_key: { "first_seen": "2026-04-23 14:00 UTC", "job": {...} } }

def extract_url_from_markdown_link(text: str) -> str:
    if not text:
        return ""
    m = re.search(r"\((https?://[^)]+)\)", text)
    return m.group(1) if m else text


def _get_url_key(job: dict) -> str:
    link = str(job.get("Link", "") or "")
    url = extract_url_from_markdown_link(link).strip()
    if url:
        return url
    return f"{job.get('Company', '')}|{job.get('Title', '')}|{job.get('Location', '')}"


def load_seen_index(path: str = "seen_jobs.json") -> dict:
    try:
        if not os.path.exists(path):
            return {}
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def save_seen_index(index: dict, path: str = "seen_jobs.json") -> None:
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(index, f, ensure_ascii=False, indent=2, sort_keys=True)
    except Exception:
        pass


def add_new_jobs_to_seen(
    jobs: list[dict],
    now: datetime,
    index_path: str = "seen_jobs.json"
) -> list[dict]:
    """
    Adds new (unseen) jobs to seen_jobs.json with full job data.
    Returns only the NEWLY added jobs from this run.
    """
    now_str = ist_timestamp_str(now)
    seen = load_seen_index(index_path)
    new_jobs = []

    for job in jobs:
        url_key = _get_url_key(job)
        if url_key not in seen:
            job["Added At"] = now_str
            seen[url_key] = {
                "first_seen": now_str,
                "job": dict(job)
            }
            new_jobs.append(job)
        else:
            entry = seen[url_key]
            # Migrate old format (string timestamp only)
            if isinstance(entry, str):
                job["Added At"] = entry
                seen[url_key] = {
                    "first_seen": entry,
                    "job": dict(job)
                }
            else:
                job["Added At"] = entry.get("first_seen", now_str)
                # Update job data in case fields changed, keep first_seen
                seen[url_key]["job"] = dict(job)

    save_seen_index(seen, index_path)
    return new_jobs


def get_todays_all_jobs(target_date: str, index_path: str = "seen_jobs.json") -> list[dict]:
    """
    Returns ALL unique jobs first seen on target_date from seen_jobs.json.
    This gives the cumulative list for the whole day.
    """
    seen = load_seen_index(index_path)
    todays_jobs = []

    for url_key, entry in seen.items():
        if isinstance(entry, dict):
            first_seen = entry.get("first_seen", "")
            if first_seen.startswith(target_date):
                job = entry.get("job", {})
                if job:
                    todays_jobs.append(job)
        # Old string format: we don't have job data, skip

    return todays_jobs


# ─────────────────────────── Chrome setup ────────────────────────────

def setup_chrome_driver() -> webdriver.Chrome:
    chrome_options = webdriver.ChromeOptions()

    chrome_bin = os.getenv("CHROME_BIN")
    if chrome_bin:
        chrome_options.binary_location = chrome_bin

    options = [
        "--window-size=1200,1200",
        "--ignore-certificate-errors",
        "--headless=new",
        "--disable-gpu",
        "--no-sandbox",
        "--disable-dev-shm-usage",
        "--disable-blink-features=AutomationControlled",
        "--lang=en-US",
    ]

    user_agents = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 13_6) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.3 Safari/605.1.15",
    ]
    chrome_options.add_argument(f"--user-agent={random.choice(user_agents)}")

    for option in options:
        chrome_options.add_argument(option)

    driver = webdriver.Chrome(options=chrome_options)
    driver.set_page_load_timeout(45)
    return driver


def polite_sleep(min_s: float = 1.5, max_s: float = 4.0) -> None:
    time.sleep(random.uniform(min_s, max_s))


def click_load_more_if_present(driver: webdriver.Chrome) -> bool:
    selectors = [
        (By.CSS_SELECTOR, "button.infinite-scroller__show-more-button"),
        (By.CSS_SELECTOR, "button.show-more-less-html__button"),
        (By.XPATH, "//button[contains(., 'Show more') or contains(., 'See more jobs')]"),
    ]
    for by, value in selectors:
        try:
            btn = WebDriverWait(driver, 2).until(EC.element_to_be_clickable((by, value)))
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", btn)
            polite_sleep(0.4, 1.0)
            btn.click()
            return True
        except Exception:
            continue
    return False


# ─────────────────────────── Scraper ─────────────────────────────────

def build_linkedin_search_url(job_title: str, location: str) -> str:
    return (
        "https://www.linkedin.com/jobs/search/?"
        "f_E=1&origin=JOB_SEARCH_PAGE_JOB_FILTER&geoId=102713980&"
        "f_TPR=r2592000&"  # Last 30 days (30d = 2592000 seconds)
        f"keywords={job_title}&location={location}&refresh=true&sortBy=DD"
    )


def is_allowed_job_title(title: str) -> bool:
    """Return True only for the requested role families and fresher variants."""
    normalized = re.sub(r"[^a-z0-9]+", " ", title.lower()).strip()
    if any(re.search(pattern, normalized) for pattern in DISALLOWED_SENIORITY_PATTERNS):
        return False
    return any(re.search(pattern, normalized) for pattern in ALLOWED_ROLE_PATTERNS)


def load_profile(path: str = "profile.json") -> dict:
    if not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as profile_file:
        data = json.load(profile_file)
    return data if isinstance(data, dict) else {}


def profile_match_score(job: dict, profile: dict) -> int:
    """Score title, skills, and experience relevance from 0 to 100."""
    if not profile or not is_allowed_job_title(job.get("Title", "")) or not is_target_location(job.get("Location", "")):
        return 0
    job_text = " ".join(str(job.get(key, "")) for key in ("Title", "Description", "Role")).lower()
    profile_text = " ".join([
        str(profile.get("resume_text", "")),
        " ".join(str(item) for item in profile.get("skills", [])),
    ]).lower()
    skill_terms = set(re.findall(r"[a-z][a-z0-9+#.\-]{2,}", profile_text))
    job_terms = set(re.findall(r"[a-z][a-z0-9+#.\-]{2,}", job_text))
    overlap = len(skill_terms & job_terms)
    skill_score = min(40, overlap * 4)
    title_score = 50 if is_allowed_job_title(job.get("Title", "")) else 0
    location_score = 10
    return min(100, title_score + skill_score + location_score)


def is_target_location(location: str) -> bool:
    """Require every saved listing to be in one of the configured locations."""
    normalized = re.sub(r"[^a-z0-9]+", " ", str(location).lower()).strip()
    return any(re.search(rf"\b{re.escape(loc)}\b", normalized) for loc in TARGET_LOCATIONS)


def _json_get(url: str) -> object:
    request = Request(url, headers={"User-Agent": "OpportunityHub/1.0"})
    with urlopen(request, timeout=20) as response:
        return json.loads(response.read().decode("utf-8"))


def _canonical_url(link: str) -> str:
    url = extract_url_from_markdown_link(link).split("?")[0].rstrip("/").lower()
    return url


def gemini_match_score(job: dict, profile: dict) -> int | None:
    """Use Gemini Flash-Lite when GEMINI_API_KEY is configured; return None on failure."""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return None
    model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite")
    prompt = {
        "job_title": job.get("Title", ""),
        "job_location": job.get("Location", ""),
        "job_description": str(job.get("Description", ""))[:6000],
        "profile_skills": profile.get("skills", []),
        "profile_experience": profile.get("experience", []),
        "instruction": "Return JSON only: {\"score\": number}. Score 0-100 for entry-level fit. Require title fit and configured location; do not invent requirements.",
    }
    body = json.dumps({"contents": [{"parts": [{"text": json.dumps(prompt)}]}], "generationConfig": {"temperature": 0, "responseMimeType": "application/json"}}).encode("utf-8")
    request = Request(
        f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={quote_plus(api_key)}",
        data=body,
        headers={"Content-Type": "application/json", "User-Agent": "OpportunityHub/1.0"},
        method="POST",
    )
    try:
        with urlopen(request, timeout=20) as response:
            result = json.loads(response.read().decode("utf-8"))
        text = result["candidates"][0]["content"]["parts"][0]["text"]
        score = json.loads(text).get("score")
        return max(0, min(100, int(score)))
    except (HTTPError, URLError, TimeoutError, KeyError, IndexError, TypeError, ValueError, json.JSONDecodeError) as error:
        logger.warning("Gemini match check failed: %s", error)
        return None


def scrape_ats_company(company: dict, requested_role: str) -> list[dict]:
    """Fetch public jobs from a company's Greenhouse, Lever, or Ashby board."""
    ats, slug = company["ats"], company["slug"]
    rows = None
    attempted = []
    for candidate_ats in dict.fromkeys([ats, "greenhouse", "lever", "ashby"]):
        attempted.append(candidate_ats)
        try:
            if candidate_ats == "greenhouse":
                payload = _json_get(f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs?content=true")
                rows = payload.get("jobs", []) if isinstance(payload, dict) else []
            elif candidate_ats == "lever":
                rows = _json_get(f"https://api.lever.co/v0/postings/{slug}?mode=json")
            else:
                payload = _json_get(f"https://api.ashbyhq.com/posting-api/job-board/{slug}")
                rows = payload.get("jobs", []) if isinstance(payload, dict) else []
            break
        except (HTTPError, URLError, TimeoutError, ValueError):
            continue
    if rows is None:
        logger.warning("No public ATS board found for %s (tried %s)", company["name"], ", ".join(attempted))
        rows = []
        return []

    jobs = []
    for row in rows:
        title = row.get("title", "")
        location = row.get("location", {})
        if isinstance(location, dict):
            location = location.get("name", "")
        description = row.get("content", row.get("descriptionPlain", row.get("description", "")))
        link = row.get("absolute_url", row.get("hostedUrl", row.get("applyUrl", "")))
        job = {"Role": requested_role.title(), "Company": company["name"], "Title": title,
               "Location": str(location), "Link": f"[Apply]({link})", "Description": BeautifulSoup(str(description), "html.parser").get_text(" ")}
        if is_allowed_job_title(title) and is_target_location(str(location)):
            jobs.append(job)
    return jobs


def scrape_public_platforms(requested_role: str, location: str) -> list[dict]:
    """Search public pages on LinkedIn, Naukri, and Internshala with Selenium."""
    searches = {
        "LinkedIn": f"https://www.linkedin.com/jobs/search/?keywords={quote_plus(requested_role)}&location={quote_plus(location)}&sortBy=DD",
        "Naukri": f"https://www.naukri.com/{quote_plus(requested_role.replace(' ', '-'))}-jobs-in-{quote_plus(location.replace(' ', '-'))}",
        "Internshala": f"https://internshala.com/jobs/{quote_plus(requested_role.replace(' ', '-'))}-jobs-in-{quote_plus(location.replace(' ', '-'))}/",
    }
    jobs = []
    for source, url in searches.items():
        try:
            driver = setup_chrome_driver()
            driver.get(url)
            polite_sleep(2.0, 4.0)
            soup = BeautifulSoup(driver.page_source or "", "html.parser")
            for link in soup.find_all("a", href=True):
                title = link.get_text(" ", strip=True)
                if is_allowed_job_title(title) and is_target_location(location) and len(title) < 150:
                    jobs.append({"Role": requested_role.title(), "Company": source, "Title": title,
                                 "Location": location, "Link": f"[Apply]({link['href']})", "Description": ""})
            driver.quit()
        except Exception as error:
            logger.warning("%s search failed: %s", source, error)
    return jobs


def scrape_all_sources(requested_role: str, location: str, profile: dict) -> list[dict]:
    jobs = []
    for company in TARGET_COMPANIES:
        jobs.extend(scrape_ats_company(company, requested_role))
    jobs.extend(scrape_public_platforms(requested_role, location))
    matching = []
    seen_links = set()
    for job in jobs:
        score = profile_match_score(job, profile)
        gemini_score = gemini_match_score(job, profile) if score >= 90 else None
        if gemini_score is not None:
            score = gemini_score
        link = _canonical_url(job.get("Link", ""))
        if score >= 90 and link and link not in seen_links:
            job["Match Score"] = score
            job.pop("Description", None)
            matching.append(job)
            seen_links.add(link)
    return matching


def scrape_linkedin_jobs(job_title: str, location: str, pages: int = 5) -> list:
    jobs = []
    driver = None
    try:
        driver = setup_chrome_driver()
        search_url = build_linkedin_search_url(job_title, location)
        logger.info("Navigating to: %s", search_url)
        driver.get(search_url)
        polite_sleep(2.0, 4.5)

        for i in range(pages):
            logger.info("Scroll batch %s/%s for '%s'", i + 1, pages, job_title)
            for _ in range(4):
                driver.execute_script("window.scrollBy(0, Math.floor(document.body.scrollHeight/4));")
                polite_sleep(0.8, 1.7)
            clicked = click_load_more_if_present(driver)
            logger.info("Load-more clicked: %s", clicked)
            polite_sleep(2.5, 5.5)

        page = driver.page_source or ""
        lowered = page.lower()
        if "unusual activity" in lowered or ("verify" in lowered and "captcha" in lowered):
            logger.warning("LinkedIn likely blocked the scrape (captcha/unusual activity).")

        soup = BeautifulSoup(page, "html.parser")
        job_listings = soup.find_all(
            "div",
            class_="base-card relative w-full hover:no-underline focus:no-underline base-card--link base-search-card base-search-card--link job-search-card",
        )
        logger.info("Found %s raw listings for '%s'", len(job_listings), job_title)

        for job in job_listings:
            try:
                title_elem = job.find("h3", class_="base-search-card__title")
                company_elem = job.find("h4", class_="base-search-card__subtitle")
                location_elem = job.find("span", class_="job-search-card__location")
                link_elem = job.find("a", class_="base-card__full-link")
                date_elem = job.find("time")

                if not all([title_elem, company_elem, location_elem, link_elem, date_elem]):
                    continue

                job_title_text = title_elem.text.strip()
                job_company = company_elem.text.strip()
                job_location = location_elem.text.strip()
                apply_link = link_elem.get("href", "#")
                date_posted = date_elem.get("datetime", "N/A")

                is_target_city = any(
                    loc in job_location.lower()
                    for loc in TARGET_LOCATIONS
                )

                if is_allowed_job_title(job_title_text) and is_target_city:
                    jobs.append({
                        "Role": job_title.title(),
                        "Company": job_company,
                        "Title": job_title_text,
                        "Location": job_location,
                        "Link": f"[Apply]({apply_link})",
                        "Date Posted": date_posted,
                    })
            except Exception:
                continue

    except Exception as e:
        logger.exception("Error during scraping: %s", str(e))
    finally:
        try:
            if driver is not None:
                driver.quit()
        except Exception:
            pass

    return jobs


# ─────────────────────────── Save helpers ────────────────────────────

def save_to_csv(data: list, filename: str) -> None:
    """Save job list to a proper CSV file (comma-separated, Excel-friendly)."""
    if not data:
        logger.warning("No data to save to CSV: %s", filename)
        return
    try:
        df = pd.DataFrame(data)
        preferred = ["Role", "Company", "Title", "Location", "Link", "Date Posted", "Match Score", "Added At"]
        cols = [c for c in preferred if c in df.columns] + [c for c in df.columns if c not in preferred]
        df = df[cols]
        df.to_csv(filename, index=False, encoding="utf-8", quoting=csv.QUOTE_ALL)
        logger.info("Saved %s jobs → %s", len(data), filename)
    except Exception as e:
        logger.exception("Error saving CSV %s: %s", filename, str(e))


def save_to_txt(data: list, filename: str) -> None:
    """Save job list to human-readable block-format TXT."""
    if not data:
        logger.warning("No data to save to TXT: %s", filename)
        return
    try:
        data = sorted(data, key=lambda x: x.get("Date Posted", ""), reverse=True)
        blocks = []
        for job in data:
            blocks.append("\n".join([
                f"ROLE        - {job.get('Role', '')}",
                f"COMPANY     - {job.get('Company', '')}",
                f"TITLE       - {job.get('Title', '')}",
                f"DATE POSTED - {job.get('Date Posted', '')}",
                f"ADDED AT    - {job.get('Added At', '')}",
                f"LOCATION    - {job.get('Location', '')}",
                f"MATCH SCORE - {job.get('Match Score', '')}%",
                f"APPLY LINK  - {job.get('Link', '')}",
            ]))
        content = "\n\n".join(blocks) + "\n"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info("Saved %s jobs → %s", len(data), filename)
    except Exception as e:
        logger.exception("Error saving TXT %s: %s", filename, str(e))


def _render_job_blocks(rows: list[dict]) -> str:
    """Render jobs in fixed block format for README."""
    blocks = []
    for r in rows:
        blocks.append("".join([
            f"ROLE        - {r.get('Role', '')}<br>\n",
            f"COMPANY     - {r.get('Company', '')}<br>\n",
            f"TITLE       - {r.get('Title', '')}<br>\n",
            f"DATE POSTED - {r.get('Date Posted', '')}<br>\n",
            f"ADDED AT    - {r.get('Added At', '')}<br>\n",
            f"LOCATION    - {r.get('Location', '')}<br>\n",
            f"MATCH SCORE - {r.get('Match Score', '')}%<br>\n",
            f"APPLY LINK  - {r.get('Link', '')}",
        ]))
    return "\n".join(blocks) + ("\n" if blocks else "")


def update_readme(all_today_jobs: list, target_date: str, now: datetime) -> None:
    """
    Update README.md with ALL cumulative jobs for today.
    Stats section shows totals; job listing section shows all jobs sorted newest first.
    """
    if not os.path.exists("README.md"):
        logger.warning("README.md not found. Skipping.")
        return

    try:
        with open("README.md", "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        # Sort jobs: newest Date Posted first, then by Added At
        jobs_sorted = sorted(
            all_today_jobs,
            key=lambda x: (x.get("Date Posted", ""), x.get("Added At", "")),
            reverse=True
        )

        df = pd.DataFrame(jobs_sorted) if jobs_sorted else pd.DataFrame()

        day_label = day_label_from_target_date(target_date)
        current_time = utc_timestamp_str(now)
        total_jobs = len(jobs_sorted)

        # ── Stats section ──
        roles_emoji = ["🔴", "🟠", "🟡", "🟢", "🔵", "🟣", "🟤", "⚫"]
        stats_md = f"## 📅 {day_label} — Live Opportunities ({total_jobs} Jobs)\n"
        stats_md += f"**Last Updated:** {current_time} | **Status:** Live ✅\n\n"
        stats_md += "### Job Categories:\n"

        if not df.empty and "Role" in df.columns:
            role_counts = df["Role"].value_counts()
            for i, (role, count) in enumerate(role_counts.items()):
                emoji = roles_emoji[i % len(roles_emoji)]
                stats_md += f"- {emoji} **{role}**: {count} jobs\n"

        stats_start = content.find("<!--START_SECTION:stats-->")
        stats_end = content.find("<!--END_SECTION:stats-->")
        if stats_start != -1 and stats_end != -1:
            content = (
                content[:stats_start + len("<!--START_SECTION:stats-->\n")]
                + stats_md
                + content[stats_end:]
            )

        # ── Jobs listing section ──
        start = content.find("<!--START_SECTION:workfetch-->")
        end = content.find("<!--END_SECTION:workfetch-->")
        if start == -1 or end == -1:
            logger.warning("Workfetch markers not found in README.md")
            return

        job_blocks = _render_job_blocks(jobs_sorted) if jobs_sorted else "No jobs found yet today.\n"

        new_content = (
            content[:start]
            + f"<!--START_SECTION:workfetch-->\n{job_blocks}"
            + content[end:]
        )

        with open("README.md", "w", encoding="utf-8") as f:
            f.write(new_content)

        logger.info("README.md updated with %s cumulative jobs for %s", total_jobs, target_date)

    except Exception as e:
        logger.exception("Error updating README: %s", str(e))


def _email_link(link: str) -> str:
    url = extract_url_from_markdown_link(link).strip()
    return url if url.startswith(("http://", "https://")) else "#"


def build_digest_html(jobs: list[dict], target_date: str, scanned_count: int) -> str:
    """Build a Gmail-friendly daily digest with one apply card per new job."""
    day_label = day_label_from_target_date(target_date)
    cards = []
    for job in sorted(jobs, key=lambda item: int(item.get("Match Score", 0)), reverse=True):
        title = html.escape(str(job.get("Title", "Untitled role")))
        company = html.escape(str(job.get("Company", "")))
        location = html.escape(str(job.get("Location", "")))
        role = html.escape(str(job.get("Role", "")))
        date_posted = html.escape(str(job.get("Date Posted", "Not specified")))
        score = html.escape(str(job.get("Match Score", "")))
        apply_url = html.escape(_email_link(str(job.get("Link", ""))), quote=True)
        cards.append(f"""
        <tr><td style="padding:0 0 18px 0;"><table width="100%" cellpadding="0" cellspacing="0"
          style="background:#171a21;border:1px solid #2b303b;border-radius:14px;color:#f4f6fb;">
          <tr><td style="padding:22px 24px 10px 24px;font-family:Arial,sans-serif;">
            <span style="font-size:20px;font-weight:700;">{title}</span>
            <span style="display:inline-block;background:#aebbd0;color:#10131a;border-radius:14px;padding:4px 9px;margin-left:8px;font-weight:700;">{score}%</span>
          </td></tr>
          <tr><td style="padding:0 24px 10px 24px;color:#aebbd0;font:14px Arial,sans-serif;">{company} &middot; {location} &middot; {role}</td></tr>
          <tr><td style="padding:0 24px 18px 24px;color:#d9deea;font:14px/1.5 Arial,sans-serif;">Posted: {date_posted}</td></tr>
          <tr><td style="padding:0 24px 22px 24px;"><a href="{apply_url}" style="background:#7898ff;color:#10131a;text-decoration:none;border-radius:10px;padding:12px 18px;font:bold 15px Arial,sans-serif;display:inline-block;">Open &amp; apply &rarr;</a></td></tr>
        </table></td></tr>""")
    if not cards:
        cards.append("<tr><td style=\"padding:28px 0;font:16px Arial,sans-serif;color:#aebbd0;\">No new matching jobs today.</td></tr>")
    return f"""<!doctype html><html><body style="margin:0;background:#0f1115;padding:24px 12px;">
    <table width="100%" cellpadding="0" cellspacing="0"><tr><td align="center">
    <table width="640" cellpadding="0" cellspacing="0" style="max-width:640px;color:#f4f6fb;">
      <tr><td style="padding:18px 0 22px;font:700 28px Arial,sans-serif;">Your job digest</td></tr>
      <tr><td style="padding:0 0 24px;color:#aebbd0;font:14px/1.6 Arial,sans-serif;">{html.escape(day_label)} &middot; scanned {scanned_count} postings &middot; {len(jobs)} new matches<br>Delhi NCR &middot; 90%+ profile match</td></tr>
      {''.join(cards)}
      <tr><td style="padding:10px 0;color:#697386;font:12px Arial,sans-serif;">Generated by OpportunityHub</td></tr>
    </table></td></tr></table></body></html>"""


def send_daily_digest(jobs: list[dict], target_date: str, scanned_count: int) -> bool:
    """Send the digest through Gmail SMTP when a Gmail app password is configured."""
    password = os.getenv("GMAIL_APP_PASSWORD")
    sender = os.getenv("GMAIL_SENDER", "ishu010.com@gmail.com")
    recipient = os.getenv("GMAIL_RECIPIENT", "ishu010.com@gmail.com")
    if not password:
        logger.warning("Digest not sent: GMAIL_APP_PASSWORD is not configured.")
        return False
    message = MIMEMultipart("alternative")
    message["Subject"] = f"{len(jobs)} jobs worth your time - {day_label_from_target_date(target_date)}"
    message["From"] = sender
    message["To"] = recipient
    message.attach(MIMEText(f"{len(jobs)} new matching jobs found. Open the HTML version to apply.", "plain", "utf-8"))
    message.attach(MIMEText(build_digest_html(jobs, target_date, scanned_count), "html", "utf-8"))
    try:
        with smtplib.SMTP("smtp.gmail.com", 587, timeout=30) as server:
            server.starttls()
            server.login(sender, password)
            server.sendmail(sender, [recipient], message.as_string())
        logger.info("Daily digest sent to %s (%s jobs)", recipient, len(jobs))
        return True
    except (OSError, smtplib.SMTPException) as error:
        logger.exception("Daily digest could not be sent: %s", error)
        return False


# ─────────────────────────── Resume profile ──────────────────────────

def extract_resume_text(resume_path: str) -> str:
    """Extract text from a PDF, DOCX, or plain-text resume."""
    extension = os.path.splitext(resume_path)[1].lower()
    if extension == ".pdf":
        from pypdf import PdfReader
        return "\n".join(page.extract_text() or "" for page in PdfReader(resume_path).pages)
    if extension == ".docx":
        from docx import Document
        document = Document(resume_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        paragraphs.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        return "\n".join(paragraphs)
    if extension in {".txt", ".md"}:
        with open(resume_path, "r", encoding="utf-8", errors="ignore") as resume_file:
            return resume_file.read()
    raise ValueError("Unsupported resume format. Use PDF, DOCX, TXT, or MD.")


def _resume_lines(text: str) -> list[str]:
    return [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if line.strip()]


def _resume_section(lines: list[str], names: set[str]) -> list[str]:
    section_lines = []
    collecting = False
    section_names = {
        "skills", "technical skills", "experience", "work experience",
        "education", "projects", "certifications", "achievements",
        "summary", "profile", "objective"
    }
    for line in lines:
        normalized = re.sub(r"[^a-z ]", "", line.lower()).strip()
        if normalized in section_names:
            collecting = normalized in names
            continue
        if collecting:
            next_normalized = re.sub(r"[^a-z ]", "", line.lower()).strip()
            if next_normalized in section_names and next_normalized not in names:
                break
            section_lines.append(line)
    return section_lines


def build_profile_from_resume(resume_path: str) -> dict:
    text = extract_resume_text(resume_path)
    lines = _resume_lines(text)
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(r"(?:\+?\d[\d ()-]{8,}\d)", text)
    links = re.findall(r"https?://[^\s)]+", text)
    name = next(
        (line for line in lines[:5] if "@" not in line and not re.search(r"\d", line)
         and len(line.split()) <= 6),
        ""
    )
    skills = _resume_section(lines, {"skills", "technical skills"})
    profile = {
        "source_file": os.path.abspath(resume_path),
        "extracted_at": ist_timestamp_str(),
        "name": name,
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0).strip() if phone_match else "",
        "links": links,
        "skills": skills,
        "experience": _resume_section(lines, {"experience", "work experience"}),
        "education": _resume_section(lines, {"education"}),
        "projects": _resume_section(lines, {"projects"}),
        "certifications": _resume_section(lines, {"certifications"}),
        "summary": _resume_section(lines, {"summary", "profile", "objective"}),
        "resume_text": text.strip(),
    }
    return profile


def create_profile_from_resume(resume_path: str, output_path: str = "profile.json") -> None:
    if not os.path.isfile(resume_path):
        raise FileNotFoundError(f"Resume not found: {resume_path}")
    profile = build_profile_from_resume(resume_path)
    with open(output_path, "w", encoding="utf-8") as profile_file:
        json.dump(profile, profile_file, ensure_ascii=False, indent=2)
    print(f"Profile created: {output_path}")


# ─────────────────────────── Main ────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Scrape strictly allowed roles or build a profile from a resume.")
    parser.add_argument("--resume", help="Path to a PDF, DOCX, TXT, or MD resume")
    parser.add_argument("--profile-output", default="profile.json", help="Output path for the extracted profile")
    args = parser.parse_args()
    if args.resume:
        try:
            create_profile_from_resume(args.resume, args.profile_output)
        except (OSError, ValueError, ImportError) as error:
            raise SystemExit(f"Could not create profile: {error}")
        raise SystemExit(0)

    now = _ist_now()          # ← IST time used everywhere
    target_date = get_target_date()
    day_label = day_label_from_target_date(target_date)
    run_time_str = ist_timestamp_str(now)  # e.g. 2026-04-23 09:56 PM IST
    time_label = now.strftime("%I:%M %p")  # e.g. 09:56 PM

    print("=" * 70)
    print(f"  OPPORTUNITYHUB — {day_label} Job Scraper")
    print(f"  Run Time (UTC) : {run_time_str}")
    print(f"  Roles          : {len(TARGET_ROLES)} roles to scrape")
    print(f"  Output         : jobs/<role>/<role>_{target_date.replace('-','_')}_{now.strftime('%H_%M')}.csv")
    print("=" * 70)
    print("  This may take 10–20 minutes (LinkedIn rate limiting)...")
    print()

    all_scraped: list[dict] = []
    location = "India"
    profile = load_profile()
    if not profile:
        raise SystemExit("profile.json not found or empty. Create it with --resume before scraping.")
    role_results: dict[str, list[dict]] = {}  # role -> jobs found this run
    run_seen_links: set[str] = set()

    # ── Step 1: Scrape each role → save to its own folder immediately ──
    for idx, role in enumerate(TARGET_ROLES, 1):
        print(f"[{idx}/{len(TARGET_ROLES)}] Scraping: '{role}'")
        print("-" * 70)

        jobs = scrape_all_sources(role, location, profile)
        unique_jobs = []
        for job in jobs:
            link_key = _canonical_url(job.get("Link", ""))
            if link_key and link_key not in run_seen_links:
                run_seen_links.add(link_key)
                unique_jobs.append(job)
        jobs = unique_jobs

        if jobs:
            # Stamp Added At for this run
            for job in jobs:
                job["Added At"] = run_time_str

            # Sort newest first
            jobs_sorted = sorted(jobs, key=lambda x: x.get("Date Posted", ""), reverse=True)

            # Save role-specific CSV + TXT
            role_csv = output_role_csv_path(role, target_date, now)
            role_txt = output_role_txt_path(role, target_date, now)
            save_to_csv(jobs_sorted, role_csv)
            save_to_txt(jobs_sorted, role_txt)

            print(f"  ✓ {len(jobs)} jobs found")
            print(f"  📁 Saved → {role_csv}")

            role_results[role] = jobs_sorted
            all_scraped.extend(jobs)
        else:
            print(f"  ✗ No matching jobs found for '{role}'")
            role_results[role] = []

        if idx < len(TARGET_ROLES):
            wait = random.randint(8, 15)
            print(f"  ⏳ Waiting {wait}s...\n")
            time.sleep(wait)

    # ── Step 2: Dedup — add new jobs to seen_jobs.json ──
    print("\n" + "=" * 70)
    print(f"  Total scraped this run : {len(all_scraped)}")
    new_jobs = add_new_jobs_to_seen(all_scraped, now)
    print(f"  🆕 New unique jobs     : {len(new_jobs)}")

    # ── Step 3: Cumulative daily README update ──
    all_today_jobs = get_todays_all_jobs(target_date)
    print(f"  📊 Total jobs today    : {len(all_today_jobs)}")
    update_readme(all_today_jobs, target_date, now)

    # ── Step 4: Email only jobs first discovered in this run ──
    send_daily_digest(new_jobs, target_date, len(all_scraped))

    # ── Step 5: Summary per role ──
    print("\n" + "=" * 70)
    print("  📋 Per-Role Summary:")
    print("-" * 70)
    for role, jobs in role_results.items():
        slug = role_to_slug(role)
        folder = f"jobs/{slug}/"
        print(f"  {'✓' if jobs else '✗'} {role.title():<35} {len(jobs):>3} jobs  →  {folder}")

    print("\n" + "=" * 70)
    print(f"  ✅ Done! {len(all_today_jobs)} total jobs for {day_label} in README.")
    print(f"  ⏰ Run completed at {run_time_str}")
    print("=" * 70)
